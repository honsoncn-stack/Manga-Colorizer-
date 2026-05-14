# -*- coding: utf-8 -*-
"""Generate customer delivery Word and PDF documents with stable Chinese fonts."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "customer_delivery"
WORD_FONT = "Microsoft YaHei"
PDF_FONT_NAME = "SimHei"
PDF_FONT_PATH = Path("C:/Windows/Fonts/simhei.ttf")


@dataclass(frozen=True)
class Section:
    title: str
    body: list[str] | None = None
    bullets: list[str] | None = None
    steps: list[str] | None = None
    table: list[tuple[str, str]] | None = None
    code: str | None = None


SETUP_DOC = {
    "title": "Manga Auto Colorizer 客户环境配置教程",
    "subtitle": "Windows / D 盘部署 / 自动上色阅读器",
    "sections": [
        Section(
            "1. 文档目的",
            body=[
                "本文用于帮助客户在一台新的 Windows 电脑上配置 Manga Auto Colorizer 的运行环境。配置完成后，客户可以通过桌面快捷方式打开软件，导入本地漫画文件，阅读黑白原图，并按页或批量生成彩图。",
                "当前交付方案默认把项目、缓存、运行目录和桌面应用都放在 D 盘，避免占用系统盘，也便于后续维护。"
            ],
        ),
        Section(
            "2. 准备条件",
            bullets=[
                "Windows 10 或 Windows 11，64 位系统。",
                "D 盘至少预留 20 GB 可用空间；如果书籍较多，建议预留更多空间。",
                "建议安装 Miniconda 或 Anaconda；如果没有安装，可把 Miniconda 安装包放入交付包。",
                "CPU 可以运行基础功能；如需更快上色，建议使用 NVIDIA 显卡并安装匹配驱动。",
                "客户只需要使用本地漫画文件，不需要浏览器插件或在线账号。"
            ],
        ),
        Section(
            "3. 交付包结构",
            body=["建议把交付包解压到 D 盘，例如 D:\\MangaColorizer_Customer_Package。交付包建议包含以下内容："],
            table=[
                ("setup_customer_environment.ps1", "自动配置环境脚本"),
                ("app\\win-unpacked\\", "Electron 桌面应用文件"),
                ("project\\", "manga-auto-colorizer 项目运行文件"),
                ("manga-colorization-v2\\", "自动上色模型代码仓库"),
                ("weights\\generator.zip", "生成器权重文件"),
                ("weights\\denoiser.pth", "线稿增强或去噪权重文件"),
                ("Miniconda3-latest-Windows-x86_64.exe", "可选的 Conda 安装包"),
                ("wheelhouse\\", "可选的离线 Python wheel 缓存")
            ],
        ),
        Section(
            "4. 一键配置步骤",
            body=["在交付包目录中打开 PowerShell，按顺序执行以下命令。脚本会创建 D 盘目录、准备 Conda 环境、复制应用文件、建立桌面快捷方式，并运行基础检查。"],
            code='Set-ExecutionPolicy -Scope Process Bypass\n.\\setup_customer_environment.ps1 -PackageRoot "D:\\MangaColorizer_Customer_Package"',
        ),
        Section(
            "5. 常用参数",
            table=[
                ("-PackageRoot", "交付包所在目录。"),
                ("-ProjectRoot", "项目安装目录，默认 D:\\AIProjects\\manga-auto-colorizer。"),
                ("-CondaEnvPath", "Conda 环境目录，默认 D:\\CondaEnvs\\manga-color-v2。"),
                ("-InstallDir", "桌面应用安装目录，默认 D:\\Programs\\Manga Auto Colorizer。"),
                ("-SkipPythonPackages", "跳过 Python 依赖安装，适合已有环境时使用。"),
                ("-SkipWeightInstall", "跳过模型权重复制，适合现场手动放置权重。"),
                ("-Wheelhouse", "指定离线 wheel 目录，适合无网络环境。")
            ],
        ),
        Section(
            "6. 配置完成后的检查",
            steps=[
                "桌面出现 Manga Auto Colorizer 快捷方式。",
                "D:\\Programs\\Manga Auto Colorizer 中存在可执行程序。",
                "D:\\AIProjects\\manga-auto-colorizer 中存在 scripts、configs、desktop 等目录。",
                "D:\\CondaEnvs\\manga-color-v2 中存在 python.exe。",
                "打开软件后，设置页显示后端状态正常。"
            ],
        ),
        Section(
            "7. 常见问题",
            table=[
                ("PowerShell 提示禁止运行脚本", "先执行 Set-ExecutionPolicy -Scope Process Bypass，再重新运行脚本。"),
                ("没有 D 盘", "当前脚本按 D 盘交付设计，请先准备 D 盘或修改脚本参数。"),
                ("模型权重缺失", "确认 generator.zip 和 denoiser.pth 已放入交付包 weights 目录。"),
                ("上色速度慢", "确认显卡驱动、CUDA 相关依赖和 PyTorch 版本匹配；CPU 模式会明显更慢。"),
                ("应用打不开", "先运行脚本的环境检查，确认后端端口未被占用，并查看 logs 目录中的运行记录。")
            ],
        ),
    ],
}


USER_DOC = {
    "title": "Manga Auto Colorizer 使用教程",
    "subtitle": "本地漫画导入、阅读、自动上色与导出完整 PDF",
    "sections": [
        Section(
            "1. 软件用途",
            body=[
                "Manga Auto Colorizer 是一款本地漫画阅读与自动上色工具。它支持导入图片文件夹、PDF 和 CBZ，把黑白漫画按页缓存到本地，并生成彩色结果。整个流程默认在本机完成，适合整理个人漫画资料或做彩页预览。",
                "软件主流程是自动上色，不包含参考图上色模式。"
            ],
        ),
        Section(
            "2. 启动方式",
            steps=[
                "双击桌面的 Manga Auto Colorizer 快捷方式。",
                "等待应用自动连接本地服务。",
                "如果页面未刷新，可点击右上角刷新按钮。"
            ],
        ),
        Section(
            "3. 导入书籍",
            steps=[
                "进入“书库”。",
                "选择本地图片文件夹、PDF 或 CBZ 文件；当前版本不支持 CBR。",
                "等待软件完成导入与分页缓存。",
                "导入完成后，书籍会显示在书库列表中。"
            ],
            bullets=[
                "图片文件夹适合已经拆页的漫画。",
                "PDF 适合扫描版或整理版漫画。",
                "CBZ 适合压缩包格式漫画。"
            ],
        ),
        Section(
            "4. 阅读器操作",
            body=["进入“阅读器”后，可以查看原图与彩图，并按需上色当前页。"],
            bullets=[
                "右方向键或 Space：下一页。",
                "左方向键：上一页。",
                "B：切换黑白原图和彩色结果。",
                "C：对当前页执行上色。",
                "鼠标放在阅读框内时，可以使用滚轮翻页。"
            ],
        ),
        Section(
            "5. 批量上色",
            steps=[
                "在阅读器中打开一本书。",
                "点击“上色后 5 页”或“整本上色”。",
                "进入“上色队列”查看进度。",
                "任务完成后回到阅读器或图库查看彩图。"
            ],
        ),
        Section(
            "6. 图库预览",
            body=["图库用于分页查看已经生成的彩图，避免一次性加载大量图片造成卡顿。鼠标放在图库区域时，可以用滚轮切换页码。"],
            bullets=[
                "可以筛选当前书籍。",
                "可以调整每页显示数量。",
                "可以打开流水线输出或书库彩页输出目录。"
            ],
        ),
        Section(
            "7. 导出 PDF",
            steps=[
                "确认目标书籍已经生成需要的彩图。",
                "在阅读器中点击导出完整 PDF。已上色页会使用彩色结果，未上色页会自动用黑白原图补齐。",
                "等待导出完成。",
                "在输出目录中打开生成的 PDF 文件。"
            ],
        ),
        Section(
            "8. 设置与维护",
            bullets=[
                "设置页可以查看项目目录、输出目录、日志目录等路径。",
                "阅读偏好可以调整阅读方向、默认缩放和滚轮翻页。",
                "维护操作可以检查环境、打开项目目录、清理临时输出或清理阅读缓存。"
            ],
        ),
        Section(
            "9. 使用建议",
            bullets=[
                "单本书页数很多时，建议先少量上色确认效果，再执行整本上色。",
                "不要把书库和输出目录放在系统盘。",
                "导入前尽量保证文件名顺序正确，避免页码错乱。",
                "运行中不要手动移动项目目录、模型目录或 Conda 环境目录。"
            ],
        ),
    ],
}


def set_run_font(run, font_name: str = WORD_FONT, size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def configure_docx_styles(doc: Document) -> None:
    styles = doc.styles
    for style_name, size, bold in [
        ("Normal", 10.5, False),
        ("Title", 20, True),
        ("Subtitle", 11, False),
        ("Heading 1", 14, True),
        ("Heading 2", 12, True),
    ]:
        style = styles[style_name]
        style.font.name = WORD_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_FONT)
        style._element.rPr.rFonts.set(qn("w:ascii"), WORD_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), WORD_FONT)
        style.font.size = Pt(size)
        style.font.bold = bold


def add_docx_paragraph(doc: Document, text: str, style: str | None = None, bold: bool = False) -> None:
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run_font(run, bold=bold)


def add_docx_code(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)


def make_docx(data: dict[str, object], output_path: Path) -> None:
    doc = Document()
    configure_docx_styles(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(str(data["title"]))
    set_run_font(title_run, size=20, bold=True)

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(str(data["subtitle"]))
    set_run_font(sub_run, size=11)

    for item in data["sections"]:
        assert isinstance(item, Section)
        add_docx_paragraph(doc, item.title, style="Heading 1", bold=True)
        for text in item.body or []:
            add_docx_paragraph(doc, text)
        for text in item.bullets or []:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(text)
            set_run_font(run)
        for idx, text in enumerate(item.steps or [], 1):
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(text)
            set_run_font(run)
        if item.table:
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            header_cells = table.rows[0].cells
            header_cells[0].text = "项目"
            header_cells[1].text = "说明"
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, bold=True)
            for key, value in item.table:
                cells = table.add_row().cells
                cells[0].text = key
                cells[1].text = value
                for cell in cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            set_run_font(run)
        if item.code:
            add_docx_code(doc, item.code)

    doc.save(output_path)


def register_pdf_font() -> None:
    if not PDF_FONT_PATH.exists():
        raise FileNotFoundError(f"Missing Chinese font: {PDF_FONT_PATH}")
    pdfmetrics.registerFont(TTFont(PDF_FONT_NAME, str(PDF_FONT_PATH)))


def pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "ChineseTitle",
            parent=base["Title"],
            fontName=PDF_FONT_NAME,
            fontSize=22,
            leading=28,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#182036"),
            wordWrap="CJK",
        ),
        "subtitle": ParagraphStyle(
            "ChineseSubtitle",
            parent=base["Normal"],
            fontName=PDF_FONT_NAME,
            fontSize=10.5,
            leading=16,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#4c5873"),
            wordWrap="CJK",
        ),
        "heading": ParagraphStyle(
            "ChineseHeading",
            parent=base["Heading1"],
            fontName=PDF_FONT_NAME,
            fontSize=13,
            leading=18,
            spaceBefore=10,
            spaceAfter=6,
            textColor=colors.HexColor("#182036"),
            wordWrap="CJK",
        ),
        "body": ParagraphStyle(
            "ChineseBody",
            parent=base["Normal"],
            fontName=PDF_FONT_NAME,
            fontSize=9.5,
            leading=15,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#20243a"),
            wordWrap="CJK",
        ),
        "code": ParagraphStyle(
            "ChineseCode",
            parent=base["Code"],
            fontName=PDF_FONT_NAME,
            fontSize=8.5,
            leading=12,
            leftIndent=8,
            backColor=colors.HexColor("#f4f4f7"),
            wordWrap="CJK",
        ),
    }


def make_pdf(data: dict[str, object], output_path: Path) -> None:
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=16 * mm,
        leftMargin=16 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
    )
    story = [
        Paragraph(str(data["title"]), styles["title"]),
        Spacer(1, 4 * mm),
        Paragraph(str(data["subtitle"]), styles["subtitle"]),
        Spacer(1, 7 * mm),
    ]

    for index, item in enumerate(data["sections"]):
        assert isinstance(item, Section)
        story.append(Paragraph(item.title, styles["heading"]))
        for text in item.body or []:
            story.append(Paragraph(text, styles["body"]))
            story.append(Spacer(1, 2 * mm))
        if item.bullets:
            for text in item.bullets:
                story.append(Paragraph(f"- {text}", styles["body"]))
            story.append(Spacer(1, 2 * mm))
        if item.steps:
            for step_index, text in enumerate(item.steps, 1):
                story.append(Paragraph(f"{step_index}. {text}", styles["body"]))
            story.append(Spacer(1, 2 * mm))
        if item.table:
            rows = [[Paragraph("项目", styles["body"]), Paragraph("说明", styles["body"])]]
            rows.extend([[Paragraph(k, styles["body"]), Paragraph(v, styles["body"])] for k, v in item.table])
            table = Table(rows, colWidths=[55 * mm, 105 * mm], hAlign="LEFT")
            table.setStyle(
                TableStyle(
                    [
                        ("FONTNAME", (0, 0), (-1, -1), PDF_FONT_NAME),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eaf3ff")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#182036")),
                        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#aab4c8")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 3 * mm))
        if item.code:
            for line in item.code.splitlines():
                story.append(Paragraph(line, styles["code"]))
            story.append(Spacer(1, 2 * mm))
        if index == 4 and data["title"].endswith("配置教程"):
            story.append(PageBreak())

    doc.build(story)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    register_pdf_font()
    make_docx(SETUP_DOC, OUT_DIR / "customer_environment_setup_guide.docx")
    make_pdf(SETUP_DOC, OUT_DIR / "customer_environment_setup_guide.pdf")
    make_docx(USER_DOC, OUT_DIR / "customer_user_guide.docx")
    make_pdf(USER_DOC, OUT_DIR / "customer_user_guide.pdf")
    print(f"Generated customer documents in {OUT_DIR}")


if __name__ == "__main__":
    main()
